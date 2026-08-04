#!/usr/bin/env python3
"""Compile the governed manuscript into submission deliverables.

Produces:
    build/the-wife-of-pontius-pilate.md    single-file Markdown
    build/the-wife-of-pontius-pilate.docx  standard manuscript format

Usage:
    python3 scripts/build_manuscript.py
    python3 scripts/build_manuscript.py --keep-section-titles
    python3 scripts/build_manuscript.py --no-docx

By default the in-chapter "## Section" headings are rendered as centered
scene breaks, which is conventional for a submitted novel. Pass
--keep-section-titles to preserve them as visible headings instead.
Source files under manuscript/ are never modified.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from bookcompile import ROOT, FRONT, BACK, chapters, matter_files, total_words  # noqa: E402

BUILD = ROOT / "build"

AUTHOR = "Dustin Ober"
SURNAME = "Ober"
TITLE = "THE WIFE OF PONTIUS PILATE"
SCENE_BREAK = "#"

# Retail-only front/back matter (copyright page, dedication, discussion
# questions, and similar apparatus) is excluded here: an agent submission
# manuscript should carry nothing but a title page, an optional epigraph
# or cast list, the text, and at most an author's note.
INCLUDE_RETAIL_ONLY = False


def build_markdown(keep_titles: bool) -> str:
    out: list[str] = []
    for _path, body in matter_files(FRONT, include_retail_only=INCLUDE_RETAIL_ONLY):
        out.append(body)
        out.append("\n---\n")
    for _, title, blocks in chapters():
        out.append(f"# {title}\n")
        for kind, text in blocks:
            if kind == "section":
                out.append(f"## {text}\n" if keep_titles else f"{SCENE_BREAK}\n")
            else:
                out.append(f"{text}\n")
        out.append("\n---\n")
    for _path, body in matter_files(BACK, include_retail_only=INCLUDE_RETAIL_ONLY):
        out.append(body)
        out.append("\n---\n")
    return "\n".join(out).rstrip() + "\n"


def add_runs(paragraph, text: str) -> None:
    """Render *italic* spans as real italics."""
    for i, chunk in enumerate(re.split(r"\*([^*]+)\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.italic = bool(i % 2)


def build_docx(keep_titles: bool, total_words: int) -> Path:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Inches, Pt

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0.5)

    section = doc.sections[0]
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))

    header_p = section.header.paragraphs[0]
    header_p.text = f"{SURNAME} / {TITLE} / "
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_number(header_p)

    def para(text="", *, align=None, indent=None, italic=False, bold=False, spaced=False):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.first_line_indent = Inches(indent if indent is not None else 0.5)
        if spaced:
            p.paragraph_format.space_before = Pt(24)
        if text:
            add_runs(p, text)
            for run in p.runs:
                if italic:
                    run.italic = True
                if bold:
                    run.bold = True
        return p

    # Title page
    para(AUTHOR, align=WD_ALIGN_PARAGRAPH.LEFT, indent=0)
    para(f"{total_words:,} words", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=0)
    for _ in range(6):
        para(indent=0)
    para(TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True)
    para("a novel", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, italic=True)
    para(f"by {AUTHOR}", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0)

    for _, title, blocks in chapters():
        doc.add_page_break()
        for _ in range(4):
            para(indent=0)
        para(title, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True)
        para(indent=0)
        first = True
        for kind, text in blocks:
            if kind == "section":
                if keep_titles:
                    para(text, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True, spaced=True)
                else:
                    para(SCENE_BREAK, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0)
                first = True
            else:
                para(text, indent=0 if first else 0.5)
                first = False

    for _path, body in matter_files(BACK, include_retail_only=INCLUDE_RETAIL_ONLY):
        doc.add_page_break()
        for line in body.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                for _ in range(3):
                    para(indent=0)
                para(line[2:], align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True)
            elif line.startswith("## "):
                para(line[3:], indent=0, bold=True, spaced=True)
            else:
                para(line, indent=0.5)

    BUILD.mkdir(exist_ok=True)
    out = BUILD / "the-wife-of-pontius-pilate.docx"
    doc.save(out)
    return out


def _add_page_number(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--keep-section-titles", action="store_true")
    ap.add_argument("--no-docx", action="store_true")
    args = ap.parse_args()

    BUILD.mkdir(exist_ok=True)
    md = build_markdown(args.keep_section_titles)
    md_path = BUILD / "the-wife-of-pontius-pilate.md"
    md_path.write_text(md, encoding="utf-8")

    total = total_words()

    print(f"chapters compiled : {len(list(chapters()))}")
    print(f"prose words       : {total:,}")
    print(f"markdown          : {md_path.relative_to(ROOT)}")
    if not args.no_docx:
        try:
            docx_path = build_docx(args.keep_section_titles, total)
            print(f"docx              : {docx_path.relative_to(ROOT)}")
        except ImportError:
            print("docx              : skipped (python-docx not installed)")


if __name__ == "__main__":
    main()
